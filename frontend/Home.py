"""
Job Intelligence AI - ChatGPT/Claude Style Interface
"""
import streamlit as st
from utils.api_client import APIClient
from utils.context_manager import ConversationContext
import io
from datetime import datetime
import os
try:
    from zoneinfo import ZoneInfo
except Exception:
    ZoneInfo = None

st.set_page_config(
    page_title="Job Intelligence AI",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed"
)

st.markdown("""
<style>
    .main .block-container {
        max-width: 800px;
        padding: 1rem 2rem 8rem 2rem;
        margin: 0 auto;
    }
    .stButton button {
        border-radius: 8px;
        border: 1px solid rgba(250, 250, 250, 0.2);
        transition: all 0.2s;
        font-size: 1.2rem;
        padding: 0.5rem;
        height: 3rem;
    }
    .stButton button:hover {
        border-color: rgba(250, 250, 250, 0.4);
    }
    [data-testid="stFileUploader"] small {
        display: none !important;
    }
    /* Align upload button with chat input */
    [data-testid="column"]:has(.stButton) {
        display: flex;
        align-items: flex-end;
        padding-bottom: 0.5rem;
    }
    /* Chat input container */
    .stChatFloatingInputContainer {
        padding-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

if 'api_client' not in st.session_state:
    st.session_state.api_client = APIClient()
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'uploaded_resume' not in st.session_state:
    st.session_state.uploaded_resume = None
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = None
if 'uploaded_resume_id' not in st.session_state:
    st.session_state.uploaded_resume_id = ''
if 'show_upload' not in st.session_state:
    st.session_state.show_upload = False
if 'resume_summary' not in st.session_state:
    st.session_state.resume_summary = None
if 'resume_matches' not in st.session_state:
    st.session_state.resume_matches = []
if 'context_manager' not in st.session_state:
    st.session_state.context_manager = ConversationContext(max_turns=10)

with st.sidebar:
    st.markdown("## 🤖 Job Intelligence AI")
    st.caption("Powered by ML agents")
    st.divider()
    
    if st.session_state.resume_text:
        st.success(f"✅ Resume Active")
        st.caption(f"📄 {st.session_state.uploaded_resume}")
        st.caption(f"🔢 {len(st.session_state.resume_text)} chars")
        # Show first 100 chars for debugging
        with st.expander("🔍 Debug"):
            st.code(st.session_state.resume_text[:100])
    else:
        st.info("💡 Upload resume")
    
    if st.session_state.messages:
        st.divider()
        st.caption(f"💬 {len(st.session_state.messages)} messages")
        
        # Show context awareness status
        summary = st.session_state.context_manager.get_summary()
        if summary['total_turns'] > 0:
            st.divider()
            st.markdown("**🧠 Context Awareness**")
            
            entities = summary['entities_tracked']
            if entities['companies'] > 0:
                st.caption(f"🏢 Tracking {entities['companies']} companies")
            if entities['locations'] > 0:
                st.caption(f"📍 Tracking {entities['locations']} locations")
            if summary['last_intent']:
                st.caption(f"💡 Last topic: {summary['last_intent'].replace('_', ' ').title()}")
            
            # Clear context button
            if st.button("🔄 Clear Context", help="Reset conversation context"):
                st.session_state.context_manager.clear_context()
                st.session_state.messages = []
                st.session_state.resume_text = None
                st.session_state.uploaded_resume = None
                st.session_state.uploaded_resume_id = ''
                st.session_state.resume_summary = None
                st.session_state.resume_matches = []  # Clear job matches
                st.session_state.resume_profile = None  # Clear profile
                st.rerun()

if not st.session_state.messages:
    # Creative time-based greetings
    tz_name = os.getenv('APP_TIMEZONE') or os.getenv('USER_TIMEZONE') or os.getenv('TZ') or 'America/New_York'
    if ZoneInfo is not None:
        try:
            hour = datetime.now(ZoneInfo(tz_name)).hour
        except Exception:
            hour = datetime.now().hour
    else:
        hour = datetime.now().hour
    
    if 2 <= hour < 5:
        greeting = "Couldn't sleep? 🌃"
        subtext = "Early morning or late night?"
    elif 5 <= hour < 7:
        greeting = "Early bird! 🌅"
        subtext = "Getting an early start!"
    elif 7 <= hour < 12:
        greeting = "Good morning! ☀️"
        subtext = "Ready to find your dream job?"
    elif 12 <= hour < 14:
        greeting = "Hope you've had lunch! 🍽️"
        subtext = "Midday job hunt?"
    elif 14 <= hour < 17:
        greeting = "Good afternoon! 🌤️"
        subtext = "Perfect time to explore opportunities"
    elif 17 <= hour < 20:
        greeting = "Good evening! 🌆"
        subtext = "Winding down with some job hunting?"
    elif 20 <= hour < 23:
        greeting = "Evening! 🌙"
        subtext = "Night time is the right time"
    else:  # 23-2 (11 PM - 2 AM)
        greeting = "Burning the midnight oil? 🔥"
        subtext = "Late night grind!"
    
    st.markdown(f"""
    <div style='text-align: center; padding: 5rem 2rem 3rem 2rem;'>
        <h2 style='font-size: 2.5rem; margin: 0; font-weight: 500;'>{greeting}</h2>
        <p style='font-size: 1.2rem; opacity: 0.7; margin: 1rem 0 0 0;'>{subtext}</p>
        <p style='font-size: 1rem; opacity: 0.6; margin: 0.5rem 0 0 0;'>Your AI job intelligence assistant</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Resume upload section - FIRST THING USER SEES
    st.markdown("---")
    with st.expander("📎 **Upload Resume (Recommended)** - Get personalized job matches", expanded=True):
        st.caption("💡 Upload your resume first for better, personalized job recommendations based on your skills and experience")
        upload_file = st.file_uploader("Choose resume (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'], key="resume_uploader_initial")
        
        if upload_file and not st.session_state.resume_text:
            current_file_id = f"{upload_file.name}_{upload_file.size}"
            
            with st.spinner("🤖 Analyzing resume with AI..."):
                try:
                    # Extract text
                    if upload_file.type == "application/pdf":
                        import PyPDF2
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(upload_file.read()))
                        resume_text = "".join([page.extract_text() for page in pdf_reader.pages])
                    elif upload_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        import docx
                        doc = docx.Document(io.BytesIO(upload_file.read()))
                        resume_text = "\n".join([para.text for para in doc.paragraphs])
                    else:
                        resume_text = upload_file.read().decode('utf-8')
                    
                    # Validate it's actually a resume
                    text_lower = resume_text.lower()
                    resume_keywords = ['experience', 'education', 'skills', 'work', 'university', 'degree', 'job', 'project']
                    has_resume_content = sum(1 for keyword in resume_keywords if keyword in text_lower) >= 3
                    
                    if not has_resume_content:
                        st.error("❌ This doesn't look like a resume. Please upload your actual resume.")
                    elif len(resume_text.strip()) >= 100:
                        # Store resume text
                        st.session_state.resume_text = resume_text
                        st.session_state.uploaded_resume_id = current_file_id
                        st.session_state.uploaded_resume = upload_file.name
                        
                        # Show success message with instructions - NO automatic job search
                        st.success(f"✅ Resume uploaded: {upload_file.name}")
                        st.info("💡 Now ask me: **'Show me jobs matching my resume'** or any other question!")
                        st.rerun()
                    else:
                        st.error("❌ Resume too short. Please upload a complete resume.")
                except Exception as e:
                    st.error(f"❌ Error reading file: {str(e)}")
    
    st.markdown("---")
    st.markdown("#### Or try these examples:")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔍 Find jobs", use_container_width=True):
            st.session_state.messages.append({"role": "user", "content": "Find me Software Engineer jobs in Boston"})
            st.session_state.trigger_query = True
            st.rerun()
        if st.button("💰 Salary info", use_container_width=True):
            st.session_state.messages.append({"role": "user", "content": "What's the average salary for Data Engineers?"})
            st.session_state.trigger_query = True
            st.rerun()
    with col2:
        if st.button("🌐 H-1B info", use_container_width=True):
            st.session_state.messages.append({"role": "user", "content": "Which companies sponsor H-1B?"})
            st.session_state.trigger_query = True
            st.rerun()
        if st.button("🎯 Career advice", use_container_width=True):
            st.session_state.messages.append({"role": "user", "content": "What skills for data science?"})
            st.session_state.trigger_query = True
            st.rerun()
else:
    col1, col2, col3 = st.columns([5, 2, 1])
    with col1:
        st.markdown("## 🤖 Job Intelligence AI")
    with col2:
        # Show resume status prominently
        if st.session_state.resume_text:
            st.success(f"✅ Resume: {st.session_state.uploaded_resume[:20]}...")
        else:
            st.info("💡 No resume uploaded")
    with col3:
        if st.button("🔄 New", use_container_width=True):
            st.session_state.messages = []
            st.session_state.uploaded_resume = None
            st.session_state.resume_text = None
            st.session_state.uploaded_resume_id = ''
            st.session_state.show_upload = False
            st.session_state.resume_summary = None
            st.session_state.resume_matches = []
            st.session_state.resume_profile = None
            st.session_state.context_manager.clear_context()
            st.rerun()
    st.divider()
    
    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            # Always show the assistant's message content
            if message.get("content"):
                st.markdown(message["content"])            
            
            # Display AI Intelligence Debug Info (if available)
            if message["role"] == "assistant" and "debug_info" in message:
                debug = message["debug_info"]
                
                with st.expander("🧠 **AI Intelligence - See How I Processed This**", expanded=False):
                    st.markdown("---")
                    
                    # Show agent and execution time
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.markdown(f"**🤖 Agent:** `{debug.get('agent_used', 'Unknown')}`")
                    with col2:
                        st.markdown(f"**⚡ {debug.get('execution_time_ms', 0)}ms**")
                    
                    st.markdown("---")
                    
                    # Step 1: Intent Analysis
                    intent_analysis = debug.get('intent_analysis', {})
                    if intent_analysis:
                        st.markdown("### 1️⃣ Understanding Your Question")
                        st.json(intent_analysis, expanded=True)
                    
                    # Step 2: SQL Generation (if applicable)
                    sql_generated = debug.get('sql_generated')
                    if sql_generated:
                        st.markdown("### 2️⃣ Generated SQL Query")
                        st.code(sql_generated, language='sql')
                    
                    # Step 3: Processing Steps
                    steps = debug.get('steps', [])
                    if steps:
                        st.markdown("### 3️⃣ Processing Steps")
                        for i, step in enumerate(steps, 1):
                            status_icon = "✅" if step.get('status') == 'complete' else "🔄"
                            st.markdown(f"{status_icon} **{step.get('step', 'Unknown')}**")
                            # Show SQL inline if present (no nested expanders)
                            if 'sql' in step:
                                st.code(step['sql'], language='sql')
                    
                    # Context awareness
                    st.markdown("---")
                    st.markdown("### 🧠 Context Awareness")
                    context_cols = st.columns(3)
                    with context_cols[0]:
                        resume_icon = "✅" if debug.get('has_resume_context') else "❌"
                        st.markdown(f"{resume_icon} **Resume**")
                    with context_cols[1]:
                        history_icon = "✅" if debug.get('has_chat_history') else "❌"
                        st.markdown(f"{history_icon} **History**")
                    with context_cols[2]:
                        st.markdown(f"💡 **{intent_analysis.get('intent', 'N/A').replace('_', ' ').title()}**")
    
    # Upload section - ALWAYS visible in the middle of conversation
    if not st.session_state.resume_text:
        with st.expander("📎 Upload Resume (Optional)", expanded=False):
            upload_file = st.file_uploader("Choose resume (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'], key="resume_uploader_inline")
            
            if upload_file:
                current_file_id = f"{upload_file.name}_{upload_file.size}"
                
                if current_file_id != st.session_state.uploaded_resume_id:
                    with st.spinner("🤖 Analyzing resume with AI..."):
                        try:
                            # Extract text
                            if upload_file.type == "application/pdf":
                                import PyPDF2
                                pdf_reader = PyPDF2.PdfReader(io.BytesIO(upload_file.read()))
                                resume_text = "".join([page.extract_text() for page in pdf_reader.pages])
                            elif upload_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                import docx
                                doc = docx.Document(io.BytesIO(upload_file.read()))
                                resume_text = "\n".join([para.text for para in doc.paragraphs])
                            else:
                                resume_text = upload_file.read().decode('utf-8')
                            
                            # Validate it's actually a resume (check for common resume keywords)
                            text_lower = resume_text.lower()
                            resume_keywords = ['experience', 'education', 'skills', 'work', 'university', 'degree', 'job', 'project']
                            has_resume_content = sum(1 for keyword in resume_keywords if keyword in text_lower) >= 3
                            
                            if not has_resume_content:
                                st.error("This doesn't look like a resume. Please upload your actual resume with work experience, education, and skills.")
                            elif len(resume_text.strip()) >= 100:
                                # Store resume text
                                st.session_state.resume_text = resume_text
                                st.session_state.uploaded_resume_id = current_file_id
                                st.session_state.uploaded_resume = upload_file.name
                                
                                # Use Cortex LLM via Agent 4 to intelligently parse resume
                                try:
                                    parse_response = st.session_state.api_client.post(
                                        '/api/resume/match',
                                        json={
                                            "resume_text": resume_text,
                                            "user_id": st.session_state.get('user_id', 'anonymous')
                                        }
                                    )
                                    
                                    if parse_response and parse_response.get('status') == 'success':
                                        profile = parse_response['profile']
                                        matches = parse_response.get('top_matches', [])
                                        
                                        # Store parsed profile AND job matches for intelligent querying
                                        st.session_state.resume_profile = profile
                                        st.session_state.resume_matches = matches  # CRITICAL: Store job matches
                                        
                                        # Create intelligent summary
                                        skills = profile.get('technical_skills', [])
                                        experience = profile.get('total_experience_years', 0)
                                        education = profile.get('education_level', 'Not specified')
                                        
                                        # Add as chat message - ONLY show analysis, NO job search
                                        summary = f"✅ **Resume uploaded: {upload_file.name}**\n\n"
                                        summary += f"**📋 Resume Analysis:**\n"
                                        summary += f"💼 Experience: {experience} years\n"
                                        summary += f"🎓 Education: {education}\n"
                                        summary += f"🔧 Key Skills: {', '.join(skills[:5]) if skills else 'Detected'}\n\n"
                                        summary += f"💡 **Ready to help!** Ask me:\n"
                                        summary += f"• 'Show me jobs matching my resume'\n"
                                        summary += f"• 'Find data engineer roles in Boston'\n"
                                        summary += f"• 'Analyze my resume'\n"
                                        summary += f"• Or any other question!"
                                        
                                        st.session_state.messages.append({
                                            "role": "assistant",
                                            "content": summary
                                        })
                                    else:
                                        # Fallback if parsing fails
                                        st.session_state.messages.append({
                                            "role": "assistant",
                                            "content": f"✅ Resume uploaded: **{upload_file.name}**\n\nLoaded into context!"
                                        })
                                except Exception as api_error:
                                    # Fallback if API fails
                                    st.session_state.messages.append({
                                        "role": "assistant",
                                        "content": f"✅ Resume uploaded: **{upload_file.name}**\n\nReady!"
                                    })
                                
                                st.rerun()
                            else:
                                st.error("Resume too short. Please upload a valid file.")
                                
                        except Exception as e:
                            st.error(f"Error: {str(e)}")


# Chat input - always visible
prompt = st.chat_input("Message Job Intelligence AI...")

if prompt:
    # Add user message to history
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    # Check if user is asking about their job matches from resume analysis
    match_keywords = ['show me the jobs', 'show jobs', 'what jobs', 'job matches', 'matched jobs', 'my matches']
    is_match_query = any(keyword in prompt.lower() for keyword in match_keywords)
    
    if is_match_query and st.session_state.get('resume_matches'):
        # Show job matches from resume analysis
        matches = st.session_state.resume_matches[:5]  # Show top 5
        
        answer = f"📋 **Top {len(matches)} Job Matches for You:**\n\n"
        for i, job in enumerate(matches, 1):
            answer += f"**{i}. {job['title']}** at {job['company']}\n"
            answer += f"   📍 {job['location']} | 🎯 Match: {job['overall_score']:.0f}%\n"
            answer += f"   💡 {job['match_reasoning']}\n"
            answer += f"   🔗 [Apply Here]({job['url']})\n\n"
        
        answer += "Ask me anything about these jobs or request more details!"
        st.session_state.messages.append({"role": "assistant", "content": answer})
        st.rerun()
    else:
        # Get AI response
        try:
            # Use intelligent context manager to enhance the query
            enhanced_prompt, metadata = st.session_state.context_manager.enhance_query(prompt)
            
            # Build chat history from context manager
            chat_history = []
            for turn in st.session_state.context_manager.conversation_history[-5:]:  # Last 5 turns
                chat_history.append({
                    "user": turn.get('user', ''),
                    "assistant": turn.get('assistant', '')[:500]  # Truncate long responses
                })
            
            # Build request with resume text AND chat history
            request_data = {
                "question": enhanced_prompt,
                "user_id": "streamlit_user",
                "chat_history": chat_history  # CRITICAL: Send conversation context
            }
            
            # Add resume text for context (persists across conversation)
            if st.session_state.resume_text:
                request_data["resume_text"] = st.session_state.resume_text
                # Debug: Show resume is being sent
                st.toast(f"📄 Context: Resume + {len(chat_history)} messages", icon="🧠")
            elif chat_history:
                st.toast(f"🧠 Sending {len(chat_history)} context messages", icon="💬")
            
            response = st.session_state.api_client.post("/api/chat/ask", json=request_data)
            
            if response and "answer" in response:
                answer = response["answer"]
                jobs = response.get("jobs", [])  # Extract jobs array from response
                
                # Update context manager with this conversation turn
                st.session_state.context_manager.update_context(prompt, answer)
                
                # Check if we have debug info to display
                debug_info = response.get("debug_info")
                
                # Build message with jobs and debug info
                message_data = {
                    "role": "assistant", 
                    "content": answer
                }
                
                # Add jobs if present
                if jobs:
                    message_data["jobs"] = jobs
                
                # Add debug info if present
                if debug_info:
                    message_data["debug_info"] = debug_info
                
                st.session_state.messages.append(message_data)
            else:
                st.session_state.messages.append({"role": "assistant", "content": "❌ Couldn't process request"})
        except Exception as e:
            st.session_state.messages.append({"role": "assistant", "content": f"❌ Error: {str(e)}"})
        
        st.rerun()
