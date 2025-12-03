"""
Resume Matcher Page
"""
import streamlit as st
from utils.api_client import APIClient

st.set_page_config(page_title="Resume Matcher", page_icon="📄", layout="wide")

# Initialize
if 'api_client' not in st.session_state:
    st.session_state.api_client = APIClient()

# Custom CSS for better styling
st.markdown("""
<style>
    .match-card {
        padding: 1.5rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
        background-color: #f8f9fa;
        margin-bottom: 1rem;
    }
    .excellent-match {
        border-left-color: #28a745;
    }
    .good-match {
        border-left-color: #ffc107;
    }
    .potential-match {
        border-left-color: #17a2b8;
    }
    .metric-container {
        background: white;
        padding: 1rem;
        border-radius: 0.5rem;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    }
    .score-badge {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 1rem;
        font-weight: 600;
        font-size: 0.875rem;
    }
    .badge-excellent {
        background-color: #d4edda;
        color: #155724;
    }
    .badge-good {
        background-color: #fff3cd;
        color: #856404;
    }
    .badge-potential {
        background-color: #d1ecf1;
        color: #0c5460;
    }
</style>
""", unsafe_allow_html=True)

st.title("📄 AI-Powered Resume Matcher")
st.markdown("Upload your resume and discover jobs tailored to your skills and experience")

# Tabs
tab1, tab2 = st.tabs(["Upload Resume", "View Matches"])

with tab1:
    st.header("📤 Upload Your Resume")
    st.markdown("Get matched with jobs that align with your experience and skills")
    
    # User info in columns
    col_u1, col_u2 = st.columns([3, 1])
    with col_u1:
        user_id = st.text_input("👤 User ID (Optional)", placeholder="Enter your user ID", help="Leave blank for guest access")
    with col_u2:
        st.markdown("&nbsp;")  # Spacing
        st.markdown("&nbsp;")
    
    st.divider()
    
    # Text input
    resume_text = st.text_area(
        "📝 Paste your resume text",
        height=250,
        placeholder="Paste your resume content here (minimum 100 characters)...",
        help="Include your skills, experience, education, and work authorization status"
    )
    
    # Or file upload
    st.markdown("**— OR —**")
    st.markdown("**📎 Upload a file:**")
    
    # Initialize resume_text if not from text area
    if not resume_text:
        resume_text = ""
    
    try:
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['txt', 'pdf', 'docx'],
            help="Supported formats: TXT, PDF, DOCX (max 200MB)",
            key="resume_uploader"
        )
    except Exception as e:
        st.error(f"File uploader error: {str(e)}")
        st.info("💡 Please paste your resume text in the box above instead")
        uploaded_file = None
    
    # Parse uploaded file
    if uploaded_file is not None:
        try:
            file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
            
            if file_size_mb > 5:
                st.warning(f"⚠️ Large file detected ({file_size_mb:.2f} MB). Processing may take a moment...")
            else:
                st.info(f"📁 File: {uploaded_file.name} ({file_size_mb:.2f} MB)")
            
            # Parse based on file type
            if uploaded_file.name.endswith('.txt'):
                try:
                    resume_text = uploaded_file.read().decode('utf-8')
                    st.success(f"✅ Text file parsed successfully - {len(resume_text)} characters")
                except Exception as e:
                    st.error(f"Error reading text file: {str(e)}")
                    resume_text = ""
            
            elif uploaded_file.name.endswith('.pdf'):
                try:
                    import PyPDF2
                    from io import BytesIO
                    
                    pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.getvalue()))
                    resume_text = ""
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            resume_text += page_text + "\n"
                    
                    if resume_text.strip():
                        st.success(f"✅ PDF parsed: {len(pdf_reader.pages)} pages, {len(resume_text)} characters")
                    else:
                        st.error("❌ Could not extract text from PDF. The file might be image-based.")
                        resume_text = ""
                        
                except ImportError:
                    st.error("❌ PDF parsing library not available. Please contact support.")
                    resume_text = ""
                except Exception as e:
                    st.error(f"❌ Error parsing PDF: {str(e)}")
                    st.info("💡 Tip: Try converting your PDF to text or using a different PDF file")
                    resume_text = ""
            
            elif uploaded_file.name.endswith('.docx'):
                try:
                    import docx
                    from io import BytesIO
                    
                    doc = docx.Document(BytesIO(uploaded_file.getvalue()))
                    resume_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    
                    if resume_text.strip():
                        st.success(f"✅ DOCX parsed: {len(doc.paragraphs)} paragraphs, {len(resume_text)} characters")
                    else:
                        st.error("❌ Document appears to be empty")
                        resume_text = ""
                        
                except ImportError:
                    st.error("❌ DOCX parsing library not available. Please contact support.")
                    resume_text = ""
                except Exception as e:
                    st.error(f"❌ Error parsing DOCX: {str(e)}")
                    resume_text = ""
        
        except Exception as e:
            st.error(f"❌ Error processing file: {str(e)}")
            st.info("💡 Try using a smaller file or paste your resume text directly")
            resume_text = ""
    
    st.divider()
    
    # Character count
    char_count = len(resume_text)
    if char_count > 0:
        if char_count < 100:
            st.warning(f"⚠️ Resume too short: {char_count}/100 characters (minimum required)")
        else:
            st.success(f"✅ Resume ready: {char_count:,} characters")
    
    col1, col2, col3 = st.columns([2, 2, 6])
    with col1:
        if st.button("🎯 Find Matches", type="primary", use_container_width=True, disabled=(char_count < 100)):
            if resume_text:
                with st.spinner("Analyzing resume and finding matches..."):
                    try:
                        response = st.session_state.api_client.post(
                            "/api/resume/match",
                            json={
                                "resume_text": resume_text,
                                "user_id": user_id if user_id else "guest"
                            }
                        )
                        
                        if not response:
                            st.error("❌ No response from server. Please check if backend is running.")
                        elif response.get('error'):
                            st.error(f"❌ Error: {response.get('detail', response.get('error'))}")
                        elif response.get('status') == 'success':
                            st.success("✅ Resume matched successfully!")
                            
                            # Show profile
                            profile = response.get('profile', {})
                            col_a, col_b, col_c = st.columns(3)
                            with col_a:
                                exp = profile.get('total_experience_years', 0)
                                st.metric("Experience", f"{exp} years")
                            with col_b:
                                skills = profile.get('technical_skills', []) + profile.get('soft_skills', [])
                                st.metric("Skills", len(skills))
                            with col_c:
                                visa = profile.get('work_authorization', 'N/A')
                                st.metric("Visa Status", visa)
                            
                            # Show education
                            if profile.get('education_level'):
                                st.info(f"🎓 Education: {profile.get('education_level')}")
                            
                            # Store matches in session
                            st.session_state.resume_matches = response.get('top_matches', [])
                            st.session_state.resume_profile = profile
                            st.session_state.total_candidates = response.get('total_candidates', 0)
                            st.session_state.resume_id = response.get('resume_id')
                            
                            # Show preview of top matches
                            st.markdown("---")
                            st.markdown("### 🎯 Top 3 Matches")
                            matches_preview = response.get('top_matches', [])[:3]
                            if matches_preview:
                                for i, match in enumerate(matches_preview, 1):
                                    score = match.get('overall_score', 0)
                                    
                                    # Determine badge style
                                    if score >= 0.8:
                                        badge_class = "badge-excellent"
                                        badge_text = "Excellent Match"
                                        emoji = "🟢"
                                    elif score >= 0.6:
                                        badge_class = "badge-good"
                                        badge_text = "Good Match"
                                        emoji = "🟡"
                                    else:
                                        badge_class = "badge-potential"
                                        badge_text = "Potential Match"
                                        emoji = "🔵"
                                    
                                    col_a, col_b = st.columns([5, 1])
                                    with col_a:
                                        st.markdown(f"{emoji} **{match.get('title')}**")
                                        st.markdown(f"*{match.get('company')}*")
                                        st.caption(f"📍 {match.get('location')} • {score:.0%} Match • {badge_text}")
                                    with col_b:
                                        if match.get('url'):
                                            st.link_button("Apply", match['url'], use_container_width=True, type="primary")
                                    
                                    if i < len(matches_preview):
                                        st.divider()
                                
                                st.success(f"✨ Found {len(st.session_state.resume_matches)} total matches! Switch to 'View Matches' tab to see all results.")
                            else:
                                st.warning("No matches found. Try updating your resume with more details.")
                        else:
                            st.error(f"Failed to match resume: {response.get('detail', 'Unknown error')}")
                    except Exception as e:
                        st.error(f"Error matching resume: {str(e)}")
            else:
                st.warning("Please provide resume text or upload a file")
    
    with col2:
        if resume_text and st.button("🔄 Clear", use_container_width=True):
            st.rerun()

with tab2:
    st.header("🎯 Your Matching Jobs")
    
    if 'resume_matches' in st.session_state and st.session_state.resume_matches:
        # Show profile summary with better design
        if 'resume_profile' in st.session_state:
            profile = st.session_state.resume_profile
            
            st.markdown("### 📋 Your Profile Summary")
            
            # Metrics row
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                exp = profile.get('total_experience_years', 0)
                st.metric("💼 Experience", f"{exp} years")
            with col2:
                st.metric("🎓 Education", profile.get('education_level', 'N/A'))
            with col3:
                st.metric("🛂 Visa Status", profile.get('work_authorization', 'N/A'))
            with col4:
                st.metric("📊 Jobs Evaluated", f"{st.session_state.get('total_candidates', 0):,}")
            
            # Skills section with better formatting
            tech_skills = profile.get('technical_skills', [])
            soft_skills = profile.get('soft_skills', [])
            
            if tech_skills or soft_skills:
                st.markdown("---")
                skill_col1, skill_col2 = st.columns(2)
                
                with skill_col1:
                    if tech_skills:
                        st.markdown("**💻 Technical Skills:**")
                        # Display as tags
                        skills_display = " • ".join(tech_skills[:20])
                        st.markdown(f"<div style='background:#f0f2f6; padding:0.75rem; border-radius:0.5rem;'>{skills_display}</div>", unsafe_allow_html=True)
                
                with skill_col2:
                    if soft_skills:
                        st.markdown("**🤝 Soft Skills:**")
                        skills_display = " • ".join(soft_skills[:15])
                        st.markdown(f"<div style='background:#f0f2f6; padding:0.75rem; border-radius:0.5rem;'>{skills_display}</div>", unsafe_allow_html=True)
            
            # Show desired roles
            if profile.get('desired_roles'):
                st.markdown("**🎯 Target Roles:** " + ", ".join(profile['desired_roles']))
        
        st.markdown("---")
        
        # Filter options
        col_f1, col_f2, col_f3 = st.columns([2, 2, 4])
        with col_f1:
            min_score = st.slider("Min Match Score", 0, 100, 0, 5, help="Filter by minimum match percentage")
        with col_f2:
            visa_filter = st.multiselect("Visa Category", 
                                        options=list(set([m.get('visa_category') for m in st.session_state.resume_matches if m.get('visa_category')])),
                                        help="Filter by visa sponsorship")
        
        # Apply filters
        filtered_matches = [
            m for m in st.session_state.resume_matches 
            if m.get('overall_score', 0) * 100 >= min_score
            and (not visa_filter or m.get('visa_category') in visa_filter)
        ]
        
        st.markdown(f"### 🎯 Showing {len(filtered_matches)} of {len(st.session_state.resume_matches)} Matches")
        
        # Display matches
        for i, match in enumerate(filtered_matches):
            score = match.get('overall_score', 0)
            
            # Color code by score
            if score >= 0.8:
                color = "🟢"
                badge = "Excellent Match"
                badge_class = "badge-excellent"
            elif score >= 0.6:
                color = "🟡"
                badge = "Good Match"
                badge_class = "badge-good"
            else:
                color = "🔵"
                badge = "Potential Match"
                badge_class = "badge-potential"
            
            with st.expander(
                f"{color} **{match.get('title')}** at {match.get('company')} — {score:.0%}",
                expanded=(i < 3)  # Auto-expand top 3
            ):
                # Header row with title and actions
                header_col1, header_col2 = st.columns([4, 1])
                with header_col1:
                    st.markdown(f"### {match.get('title')}")
                    st.markdown(f"**{match.get('company')}** • {match.get('location', 'N/A')}")
                with header_col2:
                    if match.get('url'):
                        st.link_button("🔗 Apply Now", match['url'], use_container_width=True, type="primary")
                    if st.button("💾 Save", key=f"save_{i}", use_container_width=True):
                        st.success("Saved!")
                
                st.divider()
                
                # Details section
                detail_col1, detail_col2 = st.columns([2, 1])
                
                with detail_col1:
                    if match.get('visa_category'):
                        st.markdown(f"🛂 **Visa Sponsorship:** {match.get('visa_category')}")
                    
                    if match.get('match_reasoning'):
                        st.markdown("**Why you're a great fit:**")
                        st.info(match['match_reasoning'])
                
                with detail_col2:
                    st.markdown("**Match Score Breakdown:**")
                    
                    # Score bars
                    scores = [
                        ("Overall", match.get('overall_score', 0)),
                        ("Experience", match.get('experience_score', 0)),
                        ("Visa", match.get('visa_score', 0)),
                        ("Location", match.get('location_score', 0))
                    ]
                    
                    for label, val in scores:
                        st.markdown(f"**{label}:** {val:.0%}")
                        st.progress(val if val <= 1 else val/100)
    else:
        st.info("👆 Upload and match your resume in the first tab to see results here")

# Footer with tips
st.markdown("---")
col_tip1, col_tip2, col_tip3 = st.columns(3)

with col_tip1:
    st.markdown("### 💡 Resume Tips")
    st.markdown("""
    - Use specific technical skills
    - Quantify your achievements
    - Include years of experience
    - List relevant certifications
    """)

with col_tip2:
    st.markdown("### 🎯 Better Matches")
    st.markdown("""
    - Include visa status clearly
    - Mention desired locations
    - List programming languages
    - Add industry keywords
    """)

with col_tip3:
    st.markdown("### 🔒 Privacy & Security")
    st.markdown("""
    - Data is processed securely
    - Resume not stored permanently
    - Only used for matching
    - GDPR compliant
    """)
