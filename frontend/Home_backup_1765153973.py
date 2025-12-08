"""
Job Intelligence AI - Powered by advanced ML agents
Seamlessly integrates: Job Search • Resume Matching • Salary Analytics • H-1B Intel • Career Guidance
"""
import streamlit as st
from utils.api_client import APIClient
import time
import io
import base64

# Page config
st.set_page_config(
    page_title="Job Intelligence AI",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ChatGPT-style clean interface
st.markdown("""
<style>
    /* ChatGPT-like centered layout */
    .main .block-container {
        max-width: 800px;
        padding: 2rem 1rem 8rem 1rem;
        margin: 0 auto;
    }
    
    /* Clean chat messages */
    [data-testid="stChatMessageContainer"] {
        max-width: 100%;
    }
    
    /* Professional buttons */
    .stButton button {
        border-radius: 8px;
        border: 1px solid rgba(250, 250, 250, 0.2);
        transition: all 0.2s;
        font-weight: 400;
    }
    
    .stButton button:hover {
        border-color: rgba(250, 250, 250, 0.4);
        transform: translateY(-1px);
    }
    
    /* Hide file uploader labels */
    [data-testid="stFileUploader"] > label {
        font-size: 0 !important;
    }
    
    [data-testid="stFileUploader"] small {
        display: none !important;
    }
    
    /* Clean input */
    [data-testid="stChatInput"] textarea {
        border-radius: 12px !important;
        border: 1px solid rgba(250, 250, 250, 0.15) !important;
    }
    
    /* Sidebar styling */
    [data-testid="stSidebar"] {
        background: rgba(0, 0, 0, 0.02);
        border-right: 1px solid rgba(250, 250, 250, 0.1);
    }
    
    /* Upload button next to input */
    .upload-btn {
        position: fixed;
        bottom: 24px;
        right: 24px;
        z-index: 999;
    }
</style>
""", unsafe_allow_html=True)

# Initialize API client (session-based - clears when browser closes)
if 'api_client' not in st.session_state:
    st.session_state.api_client = APIClient()

# Initialize chat history (session-based - auto clears)
if 'messages' not in st.session_state:
    st.session_state.messages = []

# Initialize uploaded resume (session-based - auto clears)
if 'uploaded_resume' not in st.session_state:
    st.session_state.uploaded_resume = None
    st.session_state.resume_text = None

# Clean header - only show when chatting
if st.session_state.messages:
    col1, col2 = st.columns([6, 1])
    with col1:
        st.markdown("## 🤖 Job Intelligence AI")
    with col2:
        if st.button("🔄 New Chat", use_container_width=True):
            st.session_state.messages = []
            st.session_state.uploaded_resume = None
            st.session_state.resume_text = None
            st.rerun()
    st.divider()

# Minimal sidebar - only essentials
with st.sidebar:
    st.markdown("## Job Intelligence AI")
    st.caption("Powered by ML agents")
    st.divider()
    
    # Resume status
    if st.session_state.resume_text:
        st.success(f"✅ Resume: {st.session_state.uploaded_resume.name if st.session_state.uploaded_resume else 'Active'}")
    else:
        st.info("💡 Upload resume for personalized matching")
    
    if st.session_state.messages:
        st.divider()
        st.caption(f"💬 {len(st.session_state.messages)} messages")
    
    st.divider()
    st.markdown("""
    **All Capabilities:**
    - Job Search
    - Salary Analytics  
    - Resume Matching
    - H-1B Sponsorship
    - Career Guidance
    """)

# Chat messages at TOP (ChatGPT style)
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Inline upload area (appears when button clicked)
if st.session_state.get('show_upload', False):
    st.markdown("---")
    upload_file = st.file_uploader("📎 Choose Resume (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'], key="inline_upload")
    
    col_a, col_b = st.columns([3, 1])
    with col_b:
        if st.button("✕ Close", use_container_width=True):
            st.session_state.show_upload = False
            st.rerun()
    
    if upload_file and upload_file != st.session_state.uploaded_resume:
        try:
            if upload_file.type == "application/pdf":
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(upload_file.read()))
                resume_text = "".join([page.extract_text() for page in pdf_reader.pages])
                st.session_state.resume_text = resume_text
            elif upload_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                import docx
                doc = docx.Document(io.BytesIO(upload_file.read()))
                resume_text = "\n".join([para.text for para in doc.paragraphs])
                st.session_state.resume_text = resume_text
            else:
                st.session_state.resume_text = upload_file.read().decode('utf-8')
            
            st.session_state.uploaded_resume = upload_file
            st.session_state.messages.append({
                "role": "assistant",
                "content": f"✅ Resume loaded! I can now provide personalized job recommendations."
            })
            st.session_state.show_upload = False
            st.rerun()
        except Exception as e:
            st.error(f"❌ {str(e)}")
    st.markdown("---")

# Chat input at BOTTOM with upload button
col1, col2 = st.columns([20, 1])
with col1:
    prompt = st.chat_input("Message Job Intelligence AI...")
with col2:
    st.markdown("<div style='padding-top: 4px;'>", unsafe_allow_html=True)
    if st.button("📎" if not st.session_state.resume_text else "✅", key="upload_btn", help="Upload Resume", use_container_width=True):
        st.session_state.show_upload = not st.session_state.get('show_upload', False)
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

if prompt:
    # Add user message to chat
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    # Display user message
    with st.chat_message("user"):
        st.markdown(prompt)
    
    # Display assistant response
    with st.chat_message("assistant"):
        message_placeholder = st.empty()
        
        # Build conversation context for the API
        # Include last 3 exchanges for context (6 messages)
        context_messages = st.session_state.messages[-7:] if len(st.session_state.messages) > 7 else st.session_state.messages
        
        # Create context string with focus on entities mentioned
        context = ""
        entities = {"location": None, "company": None, "role": None}
        
        if len(context_messages) > 1:
            # Extract entities from previous messages (both user AND assistant)
            for msg in context_messages[:-1]:
                msg_lower = msg['content'].lower()
                
                # Extract location
                if "boston" in msg_lower:
                    entities["location"] = "Boston"
                elif "seattle" in msg_lower:
                    entities["location"] = "Seattle"
                elif "new york" in msg_lower or "nyc" in msg_lower:
                    entities["location"] = "New York"
                elif "california" in msg_lower or " ca" in msg_lower:
                    entities["location"] = "California"
                
                # Extract company
                companies = ["amazon", "google", "microsoft", "meta", "apple"]
                for comp in companies:
                    if comp in msg_lower:
                        entities["company"] = comp.title()
                
                # Extract role (check both user questions AND assistant job listings)
                if "software engineer" in msg_lower or "sde" in msg_lower:
                    entities["role"] = "Software Engineer"
                elif "data engineer" in msg_lower:
                    entities["role"] = "Data Engineer"
                elif "data scientist" in msg_lower:
                    entities["role"] = "Data Scientist"
                elif "intern" in msg_lower:
                    # Check what type of intern
                    if "software" in msg_lower or "engineering" in msg_lower:
                        entities["role"] = "Software Engineering Intern"
                    elif "data" in msg_lower:
                        entities["role"] = "Data Intern"
                    else:
                        entities["role"] = "Intern"
        
        # Detect query type and enhance accordingly
        query_lower = prompt.lower()
        query_type = "chat"  # default
        full_query = prompt
        
        # ALWAYS append resume context if available (for ALL queries)
        resume_context = ""
        if st.session_state.resume_text:
            resume_context = f"\n\nMy Resume:\n{st.session_state.resume_text[:2000]}"
        
        # Check if it's a resume-related query
        resume_keywords = ["resume", "cv", "review my", "feedback", "improve", "skills", "experience", "suggest", "advice", "what should i"]
        if any(keyword in query_lower for keyword in resume_keywords):
            query_type = "resume"
            if st.session_state.resume_text:
                full_query = f"{prompt}{resume_context}"
            else:
                full_query = prompt + " (Note: No resume uploaded yet. Upload via 📎 button for personalized help.)"
        
        # Check if it's a salary query - enhance with context
        elif any(word in query_lower for word in ["salary", "pay", "earn", "compensation", "how much"]):
            query_type = "salary"
            # Always add context for salary questions
            context_parts = []
            if entities["role"]:
                # Check if role is already in the question
                if entities["role"].lower() not in query_lower:
                    full_query = f"{prompt} for {entities['role']}"
            if entities["location"] and entities["location"].lower() not in query_lower:
                full_query += f" in {entities['location']}"
        
        # Check if user wants more results from previous search
        elif any(word in query_lower for word in ["more", "additional", "other", "others", "else"]):
            query_type = "more_results"
            # Reconstruct the search query from context
            if entities["role"]:
                full_query = f"show me more {entities['role']} jobs"
                if entities["location"]:
                    full_query += f" in {entities['location']}"
            else:
                full_query = prompt  # Fallback if no context
        
        # Check if it's a job search query
        elif any(word in query_lower for word in ["find", "show", "search", "list", "jobs", "positions", "openings"]):
            query_type = "job_search"
            # Enhance with context
            if entities["role"] and entities["role"].lower() not in query_lower:
                full_query = f"{prompt} for {entities['role']}"
            if entities["location"] and entities["location"].lower() not in query_lower:
                full_query += f" in {entities['location']}"
            # Add resume context for job matching
            if resume_context:
                full_query += resume_context
        
        # Check if it's H-1B/sponsorship query
        elif any(word in query_lower for word in ["sponsor", "h-1b", "h1b", "visa", "contact", "who should"]):
            query_type = "h1b"
            if entities["company"] and entities["company"].lower() not in query_lower:
                full_query = f"{prompt} at {entities['company']}"
            elif entities["role"] and entities["role"].lower() not in query_lower:
                full_query = f"{prompt} for {entities['role']}"
        
        # If question is vague but we have context, enhance it
        elif any(prompt.lower().startswith(v) for v in ["what", "tell me", "how about", "show me"]):
            query_type = "contextual"
            context_parts = []
            if entities["role"]:
                context_parts.append(f"for {entities['role']}")
            if entities["location"]:
                context_parts.append(f"in {entities['location']}")
            if entities["company"]:
                context_parts.append(f"at {entities['company']}")
            
            if context_parts:
                full_query = f"{prompt} {' '.join(context_parts)}"
        
        try:
            # Show what we're asking (for debugging)
            with st.expander("🔍 Query Analysis", expanded=False):
                st.write("**Query Type:**", query_type)
                st.write("**Original Question:**", prompt)
                st.write("**Enhanced Query:**", full_query[:500] + "..." if len(full_query) > 500 else full_query)
                if any(entities.values()):
                    st.write("**Context:**", {k: v for k, v in entities.items() if v})
                if st.session_state.resume_text:
                    st.write("**Resume:**", f"✅ Loaded ({len(st.session_state.resume_text)} chars)")
            
            # Show typing indicator
            with st.spinner("🤔 Thinking..."):
                from urllib.parse import quote
                
                # For job search queries, also fetch jobs
                jobs_data = None
                if query_type == "job_search":
                    # Extract search parameters
                    search_params = {}
                    if entities["role"]:
                        search_params["query"] = entities["role"]
                    if entities["location"]:
                        search_params["location"] = entities["location"]
                    
                    # Try to search for jobs
                    try:
                        jobs_response = st.session_state.api_client.get("/api/search", params=search_params)
                        if jobs_response and "jobs" in jobs_response:
                            jobs_data = jobs_response["jobs"][:10]  # Top 10 jobs
                    except:
                        pass
                
                # Call the chat API for answer
                response = st.session_state.api_client.post(
                    f"/api/chat/ask?question={quote(full_query)}"
                )
            
            if response and "answer" in response:
                answer = response["answer"]
                
                # Display the answer
                message_placeholder.markdown(answer)
                
                # If we have jobs, display them
                if jobs_data and len(jobs_data) > 0:
                    st.markdown("---")
                    st.markdown(f"### 🎯 Found {len(jobs_data)} Matching Jobs")
                    
                    for idx, job in enumerate(jobs_data, 1):
                        with st.expander(f"**{idx}. {job.get('title', 'N/A')}** at {job.get('company', 'N/A')}", expanded=(idx <= 3)):
                            cols = st.columns([3, 1])
                            with cols[0]:
                                st.write(f"📍 **Location:** {job.get('location', 'Not specified')}")
                                if job.get('description'):
                                    st.write(f"📝 {job['description'][:200]}...")
                                if job.get('h1b_sponsor'):
                                    st.write("🎫 **H-1B Sponsor:** ✅ Yes")
                            with cols[1]:
                                if job.get('url'):
                                    st.link_button("Apply →", job['url'], use_container_width=True)
                                if job.get('similarity_score'):
                                    st.metric("Match", f"{job['similarity_score']:.0%}")
                
                # Store assistant message with metadata
                st.session_state.messages.append({
                    "role": "assistant",
                    "content": answer,
                    "metadata": {
                        "confidence": response.get("confidence", 0),
                        "data_points": response.get("data_points", 0),
                        "sources": response.get("sources", [])
                    }
                })
                
                # Show quick follow-up suggestions
                if response.get("confidence", 0) > 0.5:
                    st.markdown("---")
                    st.markdown("**💡 Quick follow-ups:**")
                    
                    # Generate smart follow-up suggestions based on the actual response
                    suggestions = []
                    
                    # If we showed jobs, suggest salary and H-1B info
                    if "Found" in answer and "jobs" in answer.lower():
                        if entities["role"]:
                            suggestions = [
                                f"What's the salary for {entities['role']}?",
                                f"Which companies sponsor H-1B for {entities['role']}?",
                                "Tell me about the top company"
                            ]
                        else:
                            suggestions = [
                                "What's the average salary?",
                                "Which companies sponsor H-1B?",
                                "Show me more details"
                            ]
                    
                    # If we showed salary, suggest jobs and H-1B
                    elif "salary" in answer.lower() or "average" in answer.lower():
                        if entities["role"]:
                            suggestions = [
                                f"Find me {entities['role']} jobs",
                                f"Which companies sponsor H-1B for {entities['role']}?",
                                "What about in other cities?"
                            ]
                        else:
                            suggestions = [
                                "Show me related jobs",
                                "Which companies sponsor H-1B?",
                                "Compare with other roles"
                            ]
                    
                    # If we showed H-1B info, suggest jobs and contacts
                    elif "sponsor" in answer.lower() or "h-1b" in answer.lower() or "h1b" in answer.lower():
                        suggestions = [
                            "Show me jobs at these companies",
                            "What are the salaries?",
                            "Who should I contact?"
                        ]
                    
                    # Default suggestions
                    else:
                        suggestions = [
                            "Show me related jobs",
                            "What's the salary?",
                            "Which companies sponsor H-1B?"
                        ]
                    
                    cols = st.columns(len(suggestions))
                    for i, suggestion in enumerate(suggestions):
                        if cols[i].button(suggestion, key=f"suggestion_{i}_{len(st.session_state.messages)}"):
                            # Add the suggestion as a new user message
                            st.session_state.messages.append({"role": "user", "content": suggestion})
                            st.rerun()
            
            else:
                error_msg = "Sorry, I couldn't process your request. Please try rephrasing your question."
                message_placeholder.error(error_msg)
                st.session_state.messages.append({
                    "role": "assistant",
                    "content": error_msg
                })
        
        except Exception as e:
            error_msg = f"❌ Error: {str(e)}"
            message_placeholder.error(error_msg)
            st.session_state.messages.append({
                "role": "assistant",
                "content": error_msg
            })

# Welcome screen - Only show when no messages
if not st.session_state.messages:
    from datetime import datetime
    
    # Get time-based greeting
    hour = datetime.now().hour
    if hour < 12:
        greeting = "Good Morning"
        emoji = "🌅"
    elif hour < 17:
        greeting = "Good Afternoon"
        emoji = "☀️"
    else:
        greeting = "Good Evening"
        emoji = "🌙"
    
    st.markdown(f"""
    <div style='text-align: center; padding: 5rem 2rem 3rem 2rem;'>
        <h1 style='font-size: 3.5rem; margin: 0; font-weight: 300;'>{emoji}</h1>
        <h2 style='font-size: 2.2rem; margin: 1.5rem 0 0.8rem 0; font-weight: 500;'>{greeting}</h2>
        <p style='font-size: 1.1rem; opacity: 0.7; margin: 0; font-weight: 300;'>I'm your AI job intelligence assistant</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Starter prompts
    st.markdown("<div style='max-width: 700px; margin: 2rem auto;'>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2, gap="medium")
    with col1:
        if st.button("🔍  Find Software Engineer jobs", use_container_width=True, key="ex1"):
            st.session_state.messages.append({"role": "user", "content": "Find me Software Engineer jobs in Boston"})
            st.rerun()
        if st.button("💰  Check salary ranges", use_container_width=True, key="ex2"):
            st.session_state.messages.append({"role": "user", "content": "What's the average salary for Data Engineers?"})
            st.rerun()
    
    with col2:
        if st.button("🌐  H-1B sponsorship info", use_container_width=True, key="ex3"):
            st.session_state.messages.append({"role": "user", "content": "Which tech companies sponsor H-1B?"})
            st.rerun()
        if st.button("🎯  Get career advice", use_container_width=True, key="ex4"):
            st.session_state.messages.append({"role": "user", "content": "What skills should I focus on for data science roles?"})
            st.rerun()
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Show upload dialog if triggered
    if st.session_state.get('show_upload', False):
        st.markdown("<div style='max-width: 500px; margin: 3rem auto;'>", unsafe_allow_html=True)
        uploaded = st.file_uploader("📎 Upload your resume for personalized insights", type=['pdf', 'docx', 'txt'], key="welcome_upload")
        if uploaded:
            try:
                if uploaded.type == "application/pdf":
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded.read()))
                    resume_text = "".join([page.extract_text() for page in pdf_reader.pages])
                    st.session_state.resume_text = resume_text
                elif uploaded.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    import docx
                    doc = docx.Document(io.BytesIO(uploaded.read()))
                    resume_text = "\n".join([para.text for para in doc.paragraphs])
                    st.session_state.resume_text = resume_text
                else:
                    st.session_state.resume_text = uploaded.read().decode('utf-8')
                
                st.session_state.uploaded_resume = uploaded
                st.session_state.messages.append({
                    "role": "assistant",
                    "content": f"✅ Perfect! Your resume is loaded. I can now provide personalized job matches and career guidance. What would you like to explore?"
                })
                st.session_state.show_upload = False
                st.rerun()
            except Exception as e:
                st.error(f"❌ {str(e)}")
        st.markdown("</div>", unsafe_allow_html=True)
